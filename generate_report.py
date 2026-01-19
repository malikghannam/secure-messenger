"""
سكريبت توليد تقرير مشروع التخرج
نظام مراسلة آمن مع تشفير ما بعد الكمي ومصادقة ثنائية TOTP ومشاركة ملفات آمنة
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def create_report():
    doc = Document()
    
    # إعداد RTL للمستند
    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)
    
    def add_heading_rtl(text, level=1):
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16 if level == 1 else 14)
        run.font.color.rgb = RGBColor(0, 51, 102)
        return p
    
    def add_para_rtl(text):
        p = doc.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        run.font.size = Pt(12)
        return p

    # ==========================================
    # صفحة الغلاف
    # ==========================================
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("الجمهورية العربية السورية")
    run.bold = True
    run.font.size = Pt(16)
    
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("وزارة التعليم العالي والبحث العلمي")
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("جامعة [اسم الجامعة]")
    run.bold = True
    run.font.size = Pt(18)
    
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("كلية الهندسة المعلوماتية")
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("قسم أمن المعلومات")
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("مشروع أُعدّ لنيل شهادة الهندسة في الهندسة المعلوماتية")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("اختصاص: أمن المعلومات")
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()

    # عنوان المشروع
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("نظام مراسلة آمن")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("مع تشفير ما بعد الكمي ومصادقة ثنائية TOTP ومشاركة ملفات آمنة")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Secure Messaging System with Post-Quantum Cryptography,")
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TOTP Two-Factor Authentication, and Secure File Sharing")
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # معلومات الطالب والمشرف
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("إعداد الطالب: محمد مالك غنام")
    run.bold = True
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("إشراف: د. كريستين زينية")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("العام الدراسي 2024-2025")
    run.bold = True
    run.font.size = Pt(14)
    
    doc.add_page_break()

    # ==========================================
    # الملخص العربي
    # ==========================================
    add_heading_rtl("الملخص", 1)
    
    abstract_ar = """يقدم هذا المشروع نظام مراسلة آمن يجمع بين تقنيات التشفير التقليدية والتشفير ما بعد الكمي (Post-Quantum Cryptography). يهدف النظام إلى توفير حماية شاملة للاتصالات الإلكترونية في مواجهة التهديدات الحالية والمستقبلية الناتجة عن تطور الحوسبة الكمية.

يعتمد النظام على بروتوكول PQX3DH الهجين الذي يدمج بين بروتوكول X3DH التقليدي وخوارزمية Kyber512 المقاومة للحوسبة الكمية، مما يوفر طبقة حماية مزدوجة. كما يستخدم خوارزمية Double Ratchet لضمان السرية التامة للأمام (Perfect Forward Secrecy)، بحيث لا يؤدي اختراق مفتاح واحد إلى كشف الرسائل السابقة.

يتضمن النظام آلية مصادقة ثنائية باستخدام TOTP وفق معيار RFC 6238، مما يعزز أمان الحسابات. كما يوفر ميزة مشاركة الملفات الآمنة مع سياسات أمنية متعددة تشمل: العرض لمرة واحدة، والملفات محدودة الوقت، وتحديد عدد مرات التحميل.

تم تطوير النظام باستخدام لغة Python مع إطار عمل Flask، واختباره باستخدام اختبارات الخصائص (Property-Based Testing). يحقق النظام خصائص أمنية متقدمة تشمل: السرية التامة للأمام، التعافي من الاختراق، مقاومة الحوسبة الكمية، وعمى الخادم."""
    
    add_para_rtl(abstract_ar)
    
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("الكلمات المفتاحية: ")
    run.bold = True
    run.font.size = Pt(12)
    run = p.add_run("التشفير ما بعد الكمي، Kyber، X3DH، Double Ratchet، TOTP، المصادقة الثنائية، التشفير من طرف لطرف، السرية التامة للأمام، مشاركة الملفات الآمنة.")
    run.font.size = Pt(12)
    
    doc.add_page_break()

    # ==========================================
    # الملخص الإنكليزي
    # ==========================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Abstract")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    abstract_en = """This project presents a secure messaging system that combines traditional cryptographic techniques with post-quantum cryptography. The system aims to provide comprehensive protection for electronic communications against current and future threats arising from the development of quantum computing.

The system relies on the hybrid PQX3DH protocol, which integrates the traditional X3DH protocol with the quantum-resistant Kyber512 algorithm, providing a dual layer of protection. It also uses the Double Ratchet algorithm to ensure Perfect Forward Secrecy, so that compromising a single key does not reveal previous messages.

The system includes two-factor authentication using TOTP according to RFC 6238 standard, enhancing account security. It also provides secure file sharing with multiple security policies including: view-once, time-limited files, and download count limits.

The system was developed using Python with the Flask framework and tested using property-based testing. The system achieves advanced security properties including: Perfect Forward Secrecy, Break-in Recovery, Post-Quantum Resistance, and Server Blindness."""
    
    p = doc.add_paragraph(abstract_en)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Keywords: ")
    run.bold = True
    run = p.add_run("Post-Quantum Cryptography, Kyber, X3DH, Double Ratchet, TOTP, Two-Factor Authentication, End-to-End Encryption, Perfect Forward Secrecy, Secure File Sharing.")
    
    doc.add_page_break()

    # ==========================================
    # قائمة المصطلحات
    # ==========================================
    add_heading_rtl("قائمة المصطلحات", 1)
    
    terms = [
        ("التشفير ما بعد الكمي", "Post-Quantum Cryptography", "PQC"),
        ("التشفير من طرف لطرف", "End-to-End Encryption", "E2EE"),
        ("السرية التامة للأمام", "Perfect Forward Secrecy", "PFS"),
        ("كلمة مرور لمرة واحدة معتمدة على الوقت", "Time-based One-Time Password", "TOTP"),
        ("آلية تغليف المفاتيح", "Key Encapsulation Mechanism", "KEM"),
        ("دالة اشتقاق المفاتيح", "Key Derivation Function", "KDF"),
        ("التشفير المصادق مع البيانات المرتبطة", "Authenticated Encryption with Associated Data", "AEAD"),
        ("رمز مصادقة الرسائل", "Message Authentication Code", "MAC"),
        ("التعلم مع الأخطاء المعيارية", "Module Learning With Errors", "MLWE"),
        ("خوارزمية السقاطة المزدوجة", "Double Ratchet Algorithm", "DR"),
        ("تبادل المفاتيح الموسع الثلاثي", "Extended Triple Diffie-Hellman", "X3DH"),
    ]
    
    # إنشاء جدول المصطلحات
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # رأس الجدول
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "الاختصار"
    hdr_cells[1].text = "المصطلح الإنكليزي"
    hdr_cells[2].text = "المصطلح العربي"
    
    for ar, en, abbr in terms:
        row_cells = table.add_row().cells
        row_cells[0].text = abbr
        row_cells[1].text = en
        row_cells[2].text = ar
    
    doc.add_page_break()

    # ==========================================
    # المقدمة
    # ==========================================
    add_heading_rtl("المقدمة", 1)
    
    # تمهيد حول أمن المعلومات
    add_heading_rtl("تمهيد حول أمن المعلومات", 2)
    
    intro1 = """في عصر التحول الرقمي المتسارع، أصبحت الاتصالات الإلكترونية جزءاً لا يتجزأ من الحياة اليومية للأفراد والمؤسسات على حد سواء. تعتمد الحكومات والشركات والأفراد على تطبيقات المراسلة الفورية لتبادل المعلومات الحساسة، بدءاً من البيانات المالية والطبية وصولاً إلى الأسرار التجارية والمراسلات الدبلوماسية. هذا الاعتماد المتزايد جعل من أمن هذه الاتصالات أولوية قصوى.

يواجه مجال أمن المعلومات تحديات متعددة تتراوح بين الهجمات التقليدية كالتنصت واعتراض البيانات، والهجمات المتقدمة كهجمات الرجل في المنتصف (Man-in-the-Middle) وهجمات إعادة التشغيل (Replay Attacks). تعتمد معظم أنظمة التشفير الحالية على صعوبة حل مسائل رياضية معينة مثل تحليل الأعداد الكبيرة إلى عواملها الأولية (أساس خوارزمية RSA) ومسألة اللوغاريتم المتقطع (أساس خوارزميات Diffie-Hellman وECDH)."""
    add_para_rtl(intro1)
    
    # تمهيد حول التشفير ما بعد الكمي
    add_heading_rtl("تمهيد حول التشفير ما بعد الكمي", 2)
    
    intro2 = """في عام 1994، طور عالم الرياضيات Peter Shor خوارزمية كمية قادرة على حل مسألة تحليل الأعداد الكبيرة ومسألة اللوغاريتم المتقطع في زمن متعدد الحدود (Polynomial Time) بدلاً من الزمن الأسي المطلوب على الحواسيب التقليدية [1]. هذا يعني أن حاسوباً كمياً كافي القوة سيتمكن من كسر معظم أنظمة التشفير المستخدمة حالياً في غضون ساعات أو أيام بدلاً من ملايين السنين.

رغم أن الحواسيب الكمية القادرة على تنفيذ خوارزمية Shor لم تُبنَ بعد، إلا أن التقدم في هذا المجال يتسارع. في عام 2019، أعلنت Google عن تحقيق "التفوق الكمي" (Quantum Supremacy) بحاسوب يحتوي على 53 كيوبت. وفي عام 2023، أعلنت IBM عن حاسوب كمي بأكثر من 1000 كيوبت. يتوقع الخبراء أن الحواسيب الكمية القادرة على كسر التشفير الحالي قد تظهر خلال 10-15 سنة.

يُعد هجوم "احصد الآن، فك التشفير لاحقاً" (Harvest Now, Decrypt Later - HNDL) من أخطر التهديدات الحالية. في هذا الهجوم، يقوم المهاجمون بتسجيل وتخزين الاتصالات المشفرة الحالية، ثم ينتظرون حتى تتوفر حواسيب كمية قادرة على كسر التشفير. هذا يعني أن البيانات المشفرة اليوم قد تُكشف مستقبلاً، مما يجعل الحماية الاستباقية ضرورة ملحة."""
    add_para_rtl(intro2)

    # المشكلة العلمية
    add_heading_rtl("المشكلة العلمية", 2)
    
    problem = """تتمثل المشكلة العلمية في الحاجة إلى نظام مراسلة يجمع بين عدة متطلبات أمنية متقدمة:

أولاً: حماية ضد الهجمات الكلاسيكية الحالية باستخدام خوارزميات تشفير مثبتة الأمان مثل X25519 وXChaCha20-Poly1305.

ثانياً: حماية ضد هجمات الحواسيب الكمية المستقبلية باستخدام خوارزميات ما بعد الكمي مثل Kyber512 المعتمدة من NIST.

ثالثاً: ضمان السرية التامة للأمام (Perfect Forward Secrecy) بحيث لا يؤدي اختراق مفتاح واحد إلى كشف الرسائل السابقة أو المستقبلية.

رابعاً: مصادقة قوية متعددة العوامل لحماية الحسابات من الاختراق حتى في حالة تسريب كلمة المرور.

خامساً: سياسات أمنية متقدمة لمشاركة الملفات تتيح التحكم في دورة حياة الملف بعد إرساله.

سادساً: عمى الخادم (Server Blindness) بحيث لا يستطيع الخادم قراءة محتوى الرسائل أو الملفات."""
    add_para_rtl(problem)
    
    # الهدف من البحث
    add_heading_rtl("الهدف من البحث", 2)
    
    goals = """يهدف هذا المشروع إلى تصميم وتنفيذ نظام مراسلة آمن يحقق الأهداف التالية:

أولاً: تصميم وتنفيذ بروتوكول تبادل مفاتيح هجين (PQX3DH) يجمع بين التشفير التقليدي على منحنى X25519 والتشفير ما بعد الكمي باستخدام Kyber512. هذا النهج الهجين يضمن أن النظام يبقى آمناً حتى لو كُسرت إحدى الخوارزميتين مستقبلاً.

ثانياً: تطبيق خوارزمية Double Ratchet لضمان السرية التامة للأمام والتعافي من الاختراق. كل رسالة تُشفر بمفتاح فريد يُحذف فور استخدامه، مما يحد من الضرر في حالة الاختراق.

ثالثاً: دمج نظام TOTP للمصادقة الثنائية وفق معيار RFC 6238، مع دعم تطبيقات المصادقة القياسية مثل Google Authenticator وAuthy، بالإضافة إلى رموز النسخ الاحتياطي للطوارئ.

رابعاً: تنفيذ نظام مشاركة ملفات آمن مع سياسات أمنية متعددة تشمل: العرض لمرة واحدة (View Once)، والملفات محدودة الوقت (Time Limited)، وتحديد عدد مرات المشاهدة (View Count)، وتاريخ انتهاء الصلاحية (Expiry Date).

خامساً: ضمان التشفير من طرف لطرف (End-to-End Encryption) بحيث لا يستطيع الخادم قراءة محتوى الرسائل أو الملفات. الخادم يعمل فقط كوسيط لنقل البيانات المشفرة."""
    add_para_rtl(goals)

    # الدراسة المراجعية
    add_heading_rtl("الدراسة المراجعية", 2)
    
    lit_review = """تم مراجعة عدد من الدراسات والمعايير الحديثة (2024-2025) ذات الصلة بموضوع المشروع:

1. معيار NIST FIPS 203 (2024): أصدر المعهد الوطني للمعايير والتقنية الأمريكي (NIST) معيار ML-KEM (Module-Lattice-Based Key-Encapsulation Mechanism) الذي يعتمد على خوارزمية Kyber. هذا المعيار يمثل أول معيار رسمي للتشفير ما بعد الكمي، وقد تم اعتماده في هذا المشروع [2].

2. دراسة Bos وآخرون (2018): قدمت تحليلاً شاملاً لأمان خوارزمية Kyber وأثبتت أنها تحقق أمان IND-CCA2 (Indistinguishability under Adaptive Chosen Ciphertext Attack) تحت افتراض صعوبة مشكلة MLWE [3].

3. بروتوكولات Signal (2016): طورت مؤسسة Signal بروتوكولي X3DH وDouble Ratchet اللذين أصبحا المعيار الفعلي للتشفير من طرف لطرف. تم اعتماد هذين البروتوكولين كأساس لتصميم النظام [4][5].

4. بروتوكول PQXDH من Signal (2024): أعلنت Signal عن تحديث بروتوكولها ليدعم التشفير ما بعد الكمي باستخدام Kyber. هذا التحديث كان مصدر إلهام لتصميم بروتوكول PQX3DH في هذا المشروع [6].

5. معيار RFC 6238 (2011): يحدد خوارزمية TOTP لتوليد كلمات مرور لمرة واحدة معتمدة على الوقت. تم تطبيق هذا المعيار بالكامل في نظام المصادقة الثنائية [7].

6. دراسة Cohn-Gordon وآخرون (2024): قدمت إثباتاً رسمياً لأمان بروتوكول Signal وأكدت تحقيقه للسرية التامة للأمام والتعافي من الاختراق [8].

7. تقرير Cloudflare (2025): وثق تجربة نشر التشفير ما بعد الكمي على نطاق واسع وأكد جدوى استخدام النهج الهجين [9].

8. دراسة Dmitrienko وآخرون (2014): حللت نقاط ضعف المصادقة الثنائية عبر SMS وأوصت باستخدام TOTP كبديل أكثر أماناً [10].

9. معيار RFC 4226 (2005): يحدد خوارزمية HOTP الأساسية التي يبني عليها TOTP. تم استخدامه كمرجع لتنفيذ توليد الرموز [11].

10. دراسة Reese وآخرون (2024): قارنت بين طرق المصادقة الثنائية المختلفة من حيث الأمان وسهولة الاستخدام، وأكدت أن TOTP يوفر توازناً جيداً بينهما [12]."""
    add_para_rtl(lit_review)

    # جدول الدراسات المراجعية
    add_heading_rtl("جدول ملخص الدراسات المراجعية", 2)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "العلاقة بالمشروع"
    hdr_cells[1].text = "الموضوع"
    hdr_cells[2].text = "الدراسة"
    
    studies = [
        ("اعتماد Kyber512 كخوارزمية KEM", "معيار ML-KEM للتشفير ما بعد الكمي", "NIST FIPS 203 (2024)"),
        ("إثبات أمان IND-CCA2 لـ Kyber", "تحليل أمان خوارزمية Kyber", "Bos et al. (2018)"),
        ("أساس بروتوكولات التشفير", "بروتوكولات X3DH و Double Ratchet", "Signal (2016)"),
        ("أساس تصميم PQX3DH", "بروتوكول PQXDH الهجين", "Signal (2024)"),
        ("أساس تنفيذ المصادقة الثنائية", "معيار TOTP", "RFC 6238 (2011)"),
        ("تأكيد PFS و Break-in Recovery", "إثبات رسمي لأمان Signal", "Cohn-Gordon (2024)"),
        ("تأكيد جدوى النهج الهجين", "نشر PQC على نطاق واسع", "Cloudflare (2025)"),
        ("تبرير اختيار TOTP على SMS", "تحليل أمان 2FA", "Dmitrienko (2014)"),
        ("مرجع لتنفيذ توليد الرموز", "خوارزمية HOTP الأساسية", "RFC 4226 (2005)"),
        ("تأكيد توازن الأمان وسهولة الاستخدام", "مقارنة طرق 2FA", "Reese (2024)"),
    ]
    
    for relation, topic, study in studies:
        row_cells = table.add_row().cells
        row_cells[0].text = relation
        row_cells[1].text = topic
        row_cells[2].text = study
    
    doc.add_paragraph()

    # التطبيقات العملية
    add_heading_rtl("التطبيقات العملية للمشروع", 2)
    
    applications = """يمكن استخدام النظام المطور في مجالات متعددة تتطلب مستوى عالٍ من الأمان:

أولاً: المراسلات الحكومية والدبلوماسية الحساسة، حيث تتطلب هذه المراسلات حماية طويلة الأمد ضد هجمات HNDL.

ثانياً: الاتصالات المؤسسية والتجارية السرية، خاصة في قطاعات التقنية والمالية والقانونية.

ثالثاً: حماية البيانات الطبية والمالية التي تخضع لتشريعات صارمة مثل HIPAA وGDPR.

رابعاً: الصحافة الاستقصائية وحماية المصادر، حيث يحتاج الصحفيون لقنوات اتصال آمنة مع مصادرهم.

خامساً: الاتصالات الشخصية للأفراد المهتمين بالخصوصية والأمان الرقمي."""
    add_para_rtl(applications)
    
    # التحديات
    add_heading_rtl("التحديات", 2)
    
    challenges = """واجه المشروع عدة تحديات تقنية وعملية:

أولاً: تحقيق التوازن بين الأمان والأداء، حيث أن خوارزميات ما بعد الكمي تتطلب موارد حسابية أكبر وتنتج مفاتيح ونصوص مشفرة أكبر حجماً.

ثانياً: ضمان التوافق مع المعايير الدولية، خاصة أن معايير NIST للتشفير ما بعد الكمي صدرت حديثاً في 2024.

ثالثاً: تصميم واجهة مستخدم سهلة الاستخدام تخفي تعقيد العمليات التشفيرية عن المستخدم النهائي.

رابعاً: اختبار الخصائص الأمنية بشكل شامل باستخدام اختبارات الخصائص (Property-Based Testing) لضمان صحة التنفيذ.

خامساً: إدارة دورة حياة المفاتيح بشكل آمن، بما في ذلك توليدها وتخزينها وتدويرها وحذفها."""
    add_para_rtl(challenges)
    
    # موجز المقدمة
    add_heading_rtl("موجز المقدمة", 2)
    
    intro_summary = """تناولت هذه المقدمة السياق العام لمشكلة أمن الاتصالات الإلكترونية والتهديد الذي تمثله الحوسبة الكمية. تم تحديد المشكلة العلمية المتمثلة في الحاجة لنظام مراسلة يجمع بين الحماية الكلاسيكية والحماية ما بعد الكمية. كما تم استعراض الدراسات السابقة ذات الصلة وتحديد أهداف المشروع والتطبيقات العملية والتحديات المتوقعة."""
    add_para_rtl(intro_summary)
    
    doc.add_page_break()

    # ==========================================
    # الفصل الأول: الدراسة النظرية
    # ==========================================
    add_heading_rtl("الفصل الأول: الدراسة النظرية", 1)
    
    ch1_intro = """يتناول هذا الفصل المفاهيم والنظريات الأساسية التي تم الاعتماد عليها فعلياً في بناء النظام. يركز الفصل على الخوارزميات والبروتوكولات المستخدمة في المشروع، مع شرح الأسس الرياضية والأمنية لكل منها. تم تجنب ذكر أي مفهوم لم يتم توظيفه مباشرة في التنفيذ."""
    add_para_rtl(ch1_intro)
    
    # 1-1 التشفير ما بعد الكمي
    add_heading_rtl("1-1 التشفير ما بعد الكمي (Post-Quantum Cryptography)", 2)
    
    pqc_text = """التشفير ما بعد الكمي هو فرع من علم التشفير يهدف إلى تطوير خوارزميات مقاومة للهجمات التي قد تشنها الحواسيب الكمية. تعتمد هذه الخوارزميات على مسائل رياضية يُعتقد أنها صعبة الحل حتى على الحواسيب الكمية.

تم استخدام خوارزمية Kyber512 في هذا المشروع، وهي خوارزمية تغليف مفاتيح (Key Encapsulation Mechanism - KEM) تم اختيارها من قبل NIST كمعيار للتشفير ما بعد الكمي في عام 2024 تحت اسم ML-KEM ضمن معيار FIPS 203 [2].

تعتمد Kyber على مشكلة التعلم مع الأخطاء المعيارية (Module Learning With Errors - MLWE). المعادلة الأساسية هي:

b = A·s + e (mod q)

حيث A مصفوفة عشوائية عامة، وs متجه سري، وe متجه خطأ صغير، وq معامل أولي. صعوبة المشكلة تكمن في إيجاد s من معرفة A وb فقط. هذه المشكلة تُعتبر صعبة الحل حتى على الحواسيب الكمية [3].

تم اختيار مستوى Kyber512 لأنه يوفر أمان 128-bit كلاسيكي و64-bit كمي، مع حجم مفتاح عام 800 بايت وحجم نص مشفر 768 بايت. هذا يمثل توازناً مناسباً بين الأمان والأداء لتطبيقات المراسلة."""
    add_para_rtl(pqc_text)

    # جدول مستويات Kyber
    add_heading_rtl("جدول 1-1: مستويات أمان خوارزمية Kyber", 2)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = "حجم النص المشفر"
    hdr[1].text = "حجم المفتاح العام"
    hdr[2].text = "الأمان الكمي"
    hdr[3].text = "الأمان الكلاسيكي"
    hdr[4].text = "المستوى"
    
    kyber_levels = [
        ("768 بايت", "800 بايت", "64-bit", "128-bit", "Kyber512"),
        ("1088 بايت", "1184 بايت", "96-bit", "192-bit", "Kyber768"),
        ("1568 بايت", "1568 بايت", "128-bit", "256-bit", "Kyber1024"),
    ]
    
    for ct, pk, qsec, csec, level in kyber_levels:
        row = table.add_row().cells
        row[0].text = ct
        row[1].text = pk
        row[2].text = qsec
        row[3].text = csec
        row[4].text = level
    
    doc.add_paragraph()
    
    # 1-2 بروتوكول X3DH
    add_heading_rtl("1-2 بروتوكول X3DH (Extended Triple Diffie-Hellman)", 2)
    
    x3dh_text = """X3DH هو بروتوكول تبادل مفاتيح طورته مؤسسة Signal لإنشاء سر مشترك بين طرفين بشكل غير متزامن (Asynchronous). يتميز البروتوكول بقدرته على إنشاء جلسة آمنة حتى لو كان أحد الطرفين غير متصل [4].

يستخدم البروتوكول أربعة أنواع من المفاتيح:

مفتاح الهوية (Identity Key - IK): زوج مفاتيح X25519 طويل الأمد يُنشأ مرة واحدة عند التسجيل ويمثل هوية المستخدم.

المفتاح المسبق الموقع (Signed Pre-Key - SPK): زوج مفاتيح X25519 يُجدد دورياً (عادة كل أسبوع) ويُوقع بمفتاح الهوية.

المفتاح المسبق لمرة واحدة (One-time Pre-Key - OPK): مجموعة مفاتيح X25519 تُستهلك عند الاستخدام وتوفر حماية إضافية.

المفتاح المؤقت (Ephemeral Key - EK): مفتاح X25519 يُنشأ لكل جلسة جديدة.

عند بدء محادثة، يقوم المرسل (Alice) بحساب أربع عمليات Diffie-Hellman:
DH1 = DH(IK_Alice, SPK_Bob)
DH2 = DH(EK_Alice, IK_Bob)
DH3 = DH(EK_Alice, SPK_Bob)
DH4 = DH(EK_Alice, OPK_Bob) [إذا توفر]

ثم يتم دمج نتائج هذه العمليات باستخدام دالة اشتقاق المفاتيح HKDF للحصول على السر المشترك."""
    add_para_rtl(x3dh_text)

    # جدول مفاتيح X3DH
    add_heading_rtl("جدول 1-2: المفاتيح المستخدمة في بروتوكول X3DH", 2)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = "الوصف"
    hdr[1].text = "الاسم الإنكليزي"
    hdr[2].text = "الرمز"
    
    keys = [
        ("مفتاح الهوية طويل الأمد، يُنشأ مرة واحدة عند التسجيل", "Identity Key", "IK"),
        ("مفتاح مسبق موقع، يُجدد دورياً كل أسبوع", "Signed Pre-Key", "SPK"),
        ("مفتاح مسبق لمرة واحدة، يُستهلك عند الاستخدام", "One-time Pre-Key", "OPK"),
        ("مفتاح مؤقت يُنشأ لكل جلسة جديدة", "Ephemeral Key", "EK"),
    ]
    
    for desc, name, sym in keys:
        row = table.add_row().cells
        row[0].text = desc
        row[1].text = name
        row[2].text = sym
    
    doc.add_paragraph()
    
    # 1-3 خوارزمية Double Ratchet
    add_heading_rtl("1-3 خوارزمية Double Ratchet", 2)
    
    dr_text = """Double Ratchet هي خوارزمية طورتها مؤسسة Signal تجمع بين نوعين من آليات تحديث المفاتيح لضمان أقصى درجات الأمان [5].

السقاطة غير المتماثلة (DH Ratchet): يتم تحديث المفاتيح مع كل تبادل رسائل بين الطرفين باستخدام عملية Diffie-Hellman جديدة. عندما يستقبل طرف رسالة بمفتاح DH جديد، يقوم بحساب سر DH جديد واشتقاق مفتاح جذري جديد ومفتاح سلسلة جديد، ثم يولد مفتاح DH جديد للرسائل الصادرة. هذا يوفر التعافي من الاختراق (Break-in Recovery): حتى لو اخترق المهاجم حالة الجلسة، تبقى الرسائل المستقبلية آمنة بمجرد تبادل رسالة جديدة.

السقاطة المتماثلة (Symmetric Ratchet): يتم اشتقاق مفتاح جديد لكل رسالة من مفتاح السلسلة باستخدام دالة اشتقاق المفاتيح (KDF):
(new_chain_key, message_key) = KDF(chain_key)

مفتاح الرسالة يُستخدم لتشفير رسالة واحدة فقط ثم يُحذف نهائياً. هذا يضمن السرية التامة للأمام (Perfect Forward Secrecy) على مستوى الرسالة الواحدة.

المفاتيح المستخدمة في Double Ratchet:
- المفتاح الجذري (Root Key): يُحدث مع كل DH Ratchet
- مفتاح السلسلة (Chain Key): يُحدث مع كل رسالة
- مفتاح الرسالة (Message Key): يُستخدم لتشفير رسالة واحدة ثم يُحذف نهائياً

معالجة الرسائل غير المرتبة: قد تصل الرسائل بترتيب مختلف عن ترتيب إرسالها بسبب ظروف الشبكة. للتعامل مع هذا، يتم تخزين مفاتيح الرسائل المتخطاة مؤقتاً (حتى 1000 مفتاح كحد أقصى لمنع هجمات استنزاف الذاكرة)، وعند وصول رسالة متأخرة يُستخدم المفتاح المخزن لفك تشفيرها ثم يُحذف."""
    add_para_rtl(dr_text)

    # 1-4 نظام TOTP
    add_heading_rtl("1-4 نظام TOTP للمصادقة الثنائية", 2)
    
    totp_text = """TOTP (Time-based One-Time Password) هي خوارزمية توليد كلمات مرور لمرة واحدة معتمدة على الوقت، معرفة في معيار RFC 6238 [7]. تُستخدم كعامل ثانٍ للمصادقة بعد كلمة المرور.

آلية العمل: يتم حساب الرمز باستخدام المعادلة:
TOTP = HOTP(K, T)

حيث K هو المفتاح السري المشترك (160 bits عادة)، وT هو عداد الوقت المحسوب بقسمة الوقت الحالي (Unix timestamp) على خطوة الوقت (30 ثانية افتراضياً).

خوارزمية HOTP الأساسية (RFC 4226) [11]:
1. حساب HMAC-SHA1 للمفتاح السري مع عداد الوقت
2. تطبيق Dynamic Truncation لاستخراج 4 بايتات
3. تحويل النتيجة إلى رقم عشري من 6 خانات

تم اختيار TOTP في هذا المشروع للأسباب التالية:
- لا يعتمد على الشبكة (يعمل بدون اتصال إنترنت)
- مجاني ومفتوح المصدر
- معياري ومتوافق مع تطبيقات المصادقة الشائعة
- آمن ضد هجمات SIM Swapping التي تؤثر على SMS OTP [10]

الإعدادات المستخدمة في النظام:
- طول الرمز: 6 أرقام
- خطوة الوقت: 30 ثانية
- النافذة الزمنية: ±1 فترة (للتعامل مع انحراف الساعة)
- خوارزمية التجزئة: HMAC-SHA1
- طول المفتاح السري: 160 bits (32 حرف Base32)
- عدد رموز النسخ الاحتياطي: 10
- محاولات فاشلة قبل القفل: 5
- مدة القفل: 15 دقيقة"""
    add_para_rtl(totp_text)
    
    # جدول مقارنة طرق 2FA
    add_heading_rtl("جدول 1-3: مقارنة طرق المصادقة الثنائية", 2)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = "الملاحظات"
    hdr[1].text = "يتطلب شبكة"
    hdr[2].text = "سهولة الاستخدام"
    hdr[3].text = "مستوى الأمان"
    hdr[4].text = "الطريقة"
    
    methods = [
        ("عرضة لهجمات SIM Swapping", "نعم", "عالية", "متوسط", "SMS OTP"),
        ("يتطلب تطبيق مصادقة", "لا", "عالية", "عالي", "TOTP"),
        ("تكلفة عالية", "لا", "متوسطة", "عالي جداً", "Hardware Token"),
        ("يتطلب اتصال دائم", "نعم", "عالية جداً", "عالي", "Push Notification"),
    ]
    
    for notes, network, ease, security, method in methods:
        row = table.add_row().cells
        row[0].text = notes
        row[1].text = network
        row[2].text = ease
        row[3].text = security
        row[4].text = method
    
    doc.add_paragraph()

    # 1-5 التشفير المتماثل
    add_heading_rtl("1-5 التشفير المتماثل XChaCha20-Poly1305", 2)
    
    aead_text = """XChaCha20-Poly1305 هو نظام تشفير مصادق (Authenticated Encryption with Associated Data - AEAD) يجمع بين شيفرة التدفق XChaCha20 ورمز مصادقة الرسائل Poly1305 [13].

مكونات النظام:
- XChaCha20: شيفرة تدفق مشتقة من ChaCha20 مع nonce موسع (192 bits بدلاً من 96 bits)
- Poly1305: رمز مصادقة رسائل (MAC) يوفر سلامة البيانات

مواصفات XChaCha20-Poly1305:
- حجم المفتاح: 256 bits (32 بايت)
- حجم Nonce: 192 bits (24 بايت)
- حجم Tag: 128 bits (16 بايت)

تم اختيار XChaCha20-Poly1305 للأسباب التالية:
- أداء عالي على المعالجات الحديثة (أسرع من AES-GCM على المعالجات بدون تسريع AES)
- Nonce كبير (192 bits) يسمح بتوليد nonces عشوائية بأمان دون خطر التكرار
- مقاوم لهجمات التوقيت (Timing Attacks)
- مدعوم في مكتبة libsodium/PyNaCl المستخدمة في المشروع

يُستخدم هذا النظام لتشفير:
- الرسائل بعد اشتقاق مفتاح الرسالة من Double Ratchet
- الملفات المشاركة مع مفتاح فريد لكل ملف
- البيانات الوصفية للملفات والسياسات الأمنية"""
    add_para_rtl(aead_text)
    
    # جدول مواصفات XChaCha20
    add_heading_rtl("جدول 1-4: مواصفات XChaCha20-Poly1305", 2)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = "القيمة"
    hdr[1].text = "المعامل"
    
    specs = [
        ("256 bits (32 بايت)", "حجم المفتاح"),
        ("192 bits (24 بايت)", "حجم Nonce"),
        ("128 bits (16 بايت)", "حجم Tag"),
        ("غير محدود", "الحد الأقصى للرسالة"),
    ]
    
    for val, param in specs:
        row = table.add_row().cells
        row[0].text = val
        row[1].text = param
    
    doc.add_paragraph()

    # 1-6 دالة اشتقاق المفاتيح HKDF
    add_heading_rtl("1-6 دالة اشتقاق المفاتيح HKDF", 2)
    
    hkdf_text = """HKDF (HMAC-based Key Derivation Function) هي دالة اشتقاق مفاتيح معرفة في معيار RFC 5869. تُستخدم لاشتقاق مفاتيح تشفير قوية من مادة مفتاح أولية (Input Keying Material - IKM).

تتكون HKDF من مرحلتين:
1. Extract: استخراج مفتاح عشوائي زائف (PRK) من IKM باستخدام HMAC
   PRK = HMAC-Hash(salt, IKM)

2. Expand: توسيع PRK للحصول على مفاتيح بالطول المطلوب
   OKM = HMAC-Hash(PRK, info || counter)

تُستخدم HKDF في النظام لـ:
- اشتقاق المفتاح الجذري من أسرار X3DH وKyber في بروتوكول PQX3DH
- اشتقاق مفاتيح السلسلة ومفاتيح الرسائل في Double Ratchet
- اشتقاق مفاتيح التشفير من كلمات المرور (مع salt عشوائي)

الإعدادات المستخدمة:
- خوارزمية التجزئة: SHA-256
- طول المفتاح الناتج: 32 بايت (256 bits)
- Salt: قيمة ثابتة لكل استخدام (مثل "PQX3DH-salt" لبروتوكول PQX3DH)
- Info: سياق الاستخدام (مثل "PQX3DH-root" أو "DR-rk")"""
    add_para_rtl(hkdf_text)
    
    # موجز الفصل الأول
    add_heading_rtl("موجز الفصل الأول", 2)
    
    ch1_summary = """تناول هذا الفصل الأسس النظرية للتقنيات المستخدمة في النظام. تم شرح التهديد الكمي وخوارزمية Kyber512 المقاومة له والمعتمدة من NIST، وبروتوكول X3DH لتبادل المفاتيح بشكل غير متزامن، وخوارزمية Double Ratchet للسرية التامة للأمام والتعافي من الاختراق، ونظام TOTP للمصادقة الثنائية وفق RFC 6238، ونظام التشفير المتماثل XChaCha20-Poly1305، ودالة اشتقاق المفاتيح HKDF. جميع هذه التقنيات تم توظيفها فعلياً في بناء النظام."""
    add_para_rtl(ch1_summary)
    
    doc.add_page_break()

    # ==========================================
    # الفصل الثاني: تجهيز بيئة العمل
    # ==========================================
    add_heading_rtl("الفصل الثاني: تجهيز بيئة العمل", 1)
    
    ch2_intro = """يتناول هذا الفصل الأدوات والتقنيات المستخدمة في تطوير النظام، بما في ذلك لغات البرمجة وأطر العمل والمكتبات، بالإضافة إلى بنية المشروع وآلية التطوير."""
    add_para_rtl(ch2_intro)
    
    # 2-1 لغات البرمجة
    add_heading_rtl("2-1 لغات البرمجة والأدوات", 2)
    
    lang_text = """تم اختيار لغة Python 3.10+ كلغة البرمجة الرئيسية للمشروع للأسباب التالية:
- توفر مكتبات تشفير ناضجة ومدققة أمنياً مثل cryptography وPyNaCl
- دعم مكتبة liboqs-python للتشفير ما بعد الكمي
- سهولة التطوير السريع والنمذجة الأولية
- دعم ممتاز للبرمجة غير المتزامنة
- مجتمع كبير ووثائق شاملة

تم استخدام إطار عمل Flask 3.0 لبناء خادم الترحيل (Relay Server) لأنه:
- خفيف الوزن ومرن
- سهل التوسيع والتخصيص
- يدعم RESTful APIs بشكل ممتاز
- متوافق مع SQLAlchemy لإدارة قاعدة البيانات
- يدعم WebSocket عبر Flask-SocketIO للاتصال الفوري

لواجهة المستخدم، تم استخدام:
- HTML5 لهيكل الصفحات
- CSS3 للتنسيق مع دعم RTL للعربية
- JavaScript للتفاعلية وعمليات التشفير على جانب العميل"""
    add_para_rtl(lang_text)
    
    # 2-2 المكتبات
    add_heading_rtl("2-2 المكتبات المستخدمة", 2)
    
    # جدول المكتبات
    add_heading_rtl("جدول 2-1: المكتبات الرئيسية المستخدمة في المشروع", 2)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr = table.rows[0].cells
    hdr[0].text = "الاستخدام"
    hdr[1].text = "الإصدار"
    hdr[2].text = "المكتبة"
    
    libs = [
        ("التشفير المتماثل وغير المتماثل، X25519", "41.0+", "cryptography"),
        ("XChaCha20-Poly1305 عبر libsodium", "1.5+", "PyNaCl"),
        ("خوارزمية Kyber512 للتشفير ما بعد الكمي", "0.9+", "liboqs-python"),
        ("توليد والتحقق من رموز TOTP", "2.9+", "pyotp"),
        ("توليد رموز QR لإعداد TOTP", "7.4+", "qrcode"),
        ("معالجة الصور لرموز QR", "10.0+", "Pillow"),
        ("إطار عمل الويب للخادم", "3.0+", "Flask"),
        ("إدارة قاعدة البيانات ORM", "2.0+", "Flask-SQLAlchemy"),
        ("الاتصال الفوري WebSocket", "5.3+", "Flask-SocketIO"),
        ("اختبارات الخصائص", "6.0+", "hypothesis"),
    ]
    
    for use, ver, lib in libs:
        row = table.add_row().cells
        row[0].text = use
        row[1].text = ver
        row[2].text = lib
    
    doc.add_paragraph()
