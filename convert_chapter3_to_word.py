"""
تحويل الفصل الثالث من Markdown إلى Word
Convert Chapter 3 from Markdown to Word
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re


def create_word_document():
    """إنشاء مستند Word من ملف Markdown"""
    
    # قراءة محتوى الفصل
    with open('CHAPTER_3_COMPLETE.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # إنشاء مستند Word جديد
    doc = Document()
    
    # إعداد اتجاه النص من اليمين لليسار
    sections = doc.sections
    for section in sections:
        section.page_height = Inches(11.69)  # A4
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    # معالجة المحتوى سطر بسطر
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # تخطي الأسطر الفارغة
        if not line:
            i += 1
            continue
        
        # العناوين الرئيسية (##)
        if line.startswith('## '):
            title = line.replace('## ', '')
            p = doc.add_paragraph(title)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if p.runs:
                p.runs[0].font.size = Pt(16)
                p.runs[0].font.bold = True
                p.runs[0].font.name = 'Arial'
                p.runs[0].font.rtl = True
        
        # العناوين الفرعية (###)
        elif line.startswith('### '):
            title = line.replace('### ', '')
            p = doc.add_paragraph(title)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if p.runs:
                p.runs[0].font.size = Pt(14)
                p.runs[0].font.bold = True
                p.runs[0].font.name = 'Arial'
                p.runs[0].font.rtl = True
        
        # الجداول
        elif line.startswith('**جدول') or line.startswith('| '):
            # جمع كل أسطر الجدول
            table_lines = []
            
            # إذا كان عنوان الجدول
            if line.startswith('**جدول'):
                table_title = line.replace('**', '')
                p = doc.add_paragraph(table_title)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if p.runs:
                    p.runs[0].font.bold = True
                    p.runs[0].font.rtl = True
                i += 1
                continue
            
            # جمع أسطر الجدول
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            # إنشاء الجدول
            if table_lines:
                # تحليل الجدول
                rows = []
                for tline in table_lines:
                    if '---' not in tline:  # تخطي سطر الفاصل
                        cells = [cell.strip() for cell in tline.split('|')]
                        cells = [c for c in cells if c]  # إزالة الخلايا الفارغة
                        if cells:
                            rows.append(cells)
                
                if rows and len(rows) > 0:
                    # إنشاء جدول Word
                    max_cols = max(len(row) for row in rows)
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    # ملء الجدول
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data
                                
                                # تنسيق الخلية
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                    for run in paragraph.runs:
                                        run.font.name = 'Arial'
                                        run.font.size = Pt(11)
                                        run.font.rtl = True
                                        
                                        # تنسيق الصف الأول (العناوين)
                                        if row_idx == 0:
                                            run.font.bold = True
                    
                    doc.add_paragraph()  # سطر فارغ بعد الجدول
            
            continue
        
        # الكود
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if p.runs:
                    p.runs[0].font.name = 'Courier New'
                    p.runs[0].font.size = Pt(10)
                
                # خلفية رمادية للكود
                p.style = 'No Spacing'
        
        # النقاط والقوائم
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if p.runs:
                p.runs[0].font.rtl = True
        
        # الفقرات العادية
        else:
            # تخطي الخطوط الأفقية
            if line == '---':
                doc.add_paragraph()
            # تخطي أماكن المخططات
            elif '[يُدرج هنا مخطط' in line:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    p.runs[0].font.italic = True
                    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
            else:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if p.runs:
                    p.runs[0].font.name = 'Arial'
                    p.runs[0].font.size = Pt(12)
                    p.runs[0].font.rtl = True
        
        i += 1
    
    # حفظ المستند
    output_file = 'CHAPTER_3_FORMATTED.docx'
    doc.save(output_file)
    print(f"✅ تم إنشاء الملف: {output_file}")
    print(f"✅ File created: {output_file}")
    return output_file


if __name__ == '__main__':
    try:
        output = create_word_document()
        print(f"\n📄 افتح الملف: {output}")
        print(f"📄 Open file: {output}")
    except Exception as e:
        print(f"❌ خطأ: {e}")
        print(f"❌ Error: {e}")
