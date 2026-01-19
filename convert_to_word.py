"""
تحويل تقرير التشفير من Markdown إلى Word
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

def create_report():
    doc = Document()
    
    # العنوان الرئيسي
    title = doc.add_heading('تقرير اختبار نظام التشفير', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Cryptographic System Test Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # معلومات التقرير
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('التاريخ: ').bold = True
    info.add_run('يناير 2026\n')
    info.add_run('المشروع: ').bold = True
    info.add_run('Secure Messenger\n')
    info.add_run('الإصدار: ').bold = True
    info.add_run('1.0')
    
    doc.add_paragraph('─' * 50)
    
    # ملخص تنفيذي
    doc.add_heading('ملخص تنفيذي', level=1)
    doc.add_paragraph(
        'تم إجراء 26 اختبار شامل على نظام التشفير في تطبيق المراسلة الآمنة. '
        'جميع الاختبارات نجحت بنسبة 100%، مما يثبت فعالية وأمان نظام التشفير المستخدم.'
    )
    
    # جدول النتائج
    doc.add_heading('نتائج الاختبارات', level=2)
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    
    # رؤوس الجدول
    headers = ['الفئة', 'عدد الاختبارات', 'النتيجة', 'النسبة']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    # بيانات الجدول
    data = [
        ['XChaCha20-Poly1305', '8', '✅ نجاح', '100%'],
        ['تشفير الملفات', '6', '✅ نجاح', '100%'],
        ['PQ-X3DH', '3', '✅ نجاح', '100%'],
        ['Double Ratchet', '6', '✅ نجاح', '100%'],
        ['خصائص الأمان', '3', '✅ نجاح', '100%'],
        ['المجموع', '26', '✅ نجاح', '100%'],
    ]
    
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # القسم 1: XChaCha20-Poly1305
    doc.add_heading('1. اختبارات XChaCha20-Poly1305 (التشفير المتماثل)', level=1)
    
    doc.add_heading('1.1 الخوارزمية المستخدمة', level=2)
    algo_info = doc.add_paragraph()
    algo_info.add_run('• الاسم: ').bold = True
    algo_info.add_run('XChaCha20-Poly1305 AEAD\n')
    algo_info.add_run('• حجم المفتاح: ').bold = True
    algo_info.add_run('256 بت (32 بايت)\n')
    algo_info.add_run('• حجم Nonce: ').bold = True
    algo_info.add_run('192 بت (24 بايت)\n')
    algo_info.add_run('• حجم Tag: ').bold = True
    algo_info.add_run('128 بت (16 بايت)')
    
    doc.add_heading('1.2 الاختبارات المنفذة', level=2)
    
    # اختبار 1
    doc.add_heading('اختبار 1.2.1: Round-Trip (الذهاب والإياب)', level=3)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التحقق من أن التشفير ثم فك التشفير يعيد البيانات الأصلية\n')
    p.add_run('الطريقة:\n').bold = True
    p.add_run('  1. توليد مفتاح عشوائي (32 بايت)\n')
    p.add_run('  2. تشفير نص يحتوي على أحرف عربية وإنجليزية\n')
    p.add_run('  3. فك التشفير باستخدام نفس المفتاح\n')
    p.add_run('  4. مقارنة النتيجة مع النص الأصلي\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - النص المسترجع مطابق للأصلي')
    
    # اختبار 2
    doc.add_heading('اختبار 1.2.2: Property-Based Testing', level=3)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التحقق من صحة التشفير لأي بيانات عشوائية\n')
    p.add_run('الطريقة: ').bold = True
    p.add_run('استخدام مكتبة Hypothesis لتوليد 50 حالة اختبار عشوائية\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - جميع الحالات العشوائية نجحت')
    
    # اختبار 3
    doc.add_heading('اختبار 1.2.3: اختلاف النص المشفر عن الأصلي', level=3)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التأكد من أن النص المشفر مختلف تماماً عن النص الأصلي\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح\n')
    p.add_run('  • النص الأصلي: 14 بايت\n')
    p.add_run('  • النص المشفر: 54 بايت (يشمل nonce + ciphertext + tag)')
    
    # اختبار 4
    doc.add_heading('اختبار 1.2.4: مفاتيح مختلفة تنتج نتائج مختلفة', level=3)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التحقق من أن نفس النص يُشفر بشكل مختلف بمفاتيح مختلفة\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - كل مفتاح ينتج نص مشفر فريد')
    
    # اختبار 5
    doc.add_heading('اختبار 1.2.5: رفض المفتاح الخاطئ', level=3)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التأكد من فشل فك التشفير بمفتاح خاطئ\n')
    p.add_run('الطريقة:\n').bold = True
    p.add_run('  1. تشفير بمفتاح A\n')
    p.add_run('  2. محاولة فك التشفير بمفتاح B\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - تم رفض المفتاح الخاطئ مع رسالة خطأ')
    
    # اختبار 6
    doc.add_heading('اختبار 1.2.6: كشف التلاعب بالبيانات', level=3)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التحقق من كشف أي تعديل على النص المشفر\n')
    p.add_run('الطريقة:\n').bold = True
    p.add_run('  1. تشفير النص\n')
    p.add_run('  2. تعديل بايت واحد في النص المشفر\n')
    p.add_run('  3. محاولة فك التشفير\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - تم كشف التلاعب ورفض البيانات\n')
    p.add_run('الأهمية: ').bold = True
    p.add_run('يمنع هجمات Man-in-the-Middle')
    
    # اختبار 7
    doc.add_heading('اختبار 1.2.7: مصادقة البيانات الإضافية (AAD)', level=3)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التحقق من حماية البيانات الوصفية\n')
    p.add_run('الطريقة:\n').bold = True
    p.add_run('  1. تشفير مع AAD = "metadata:user123"\n')
    p.add_run('  2. فك التشفير مع نفس AAD → نجاح\n')
    p.add_run('  3. فك التشفير مع AAD مختلف → فشل\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - AAD يوفر حماية إضافية')
    
    # اختبار 8
    doc.add_heading('اختبار 1.2.8: رفض أحجام المفاتيح الخاطئة', level=3)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التأكد من قبول مفاتيح 32 بايت فقط\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - المفاتيح القصيرة والطويلة تُرفض')
    
    doc.add_page_break()
    
    # القسم 2: تشفير الملفات
    doc.add_heading('2. اختبارات تشفير الملفات', level=1)
    
    doc.add_heading('2.1 الاختبارات المنفذة', level=2)
    
    tests_2 = [
        ('اختبار 2.1.1: توليد مفاتيح الملفات', 
         'التحقق من توليد مفاتيح بالحجم الصحيح',
         '✅ نجاح - المفتاح 32 بايت'),
        ('اختبار 2.1.2: تفرد المفاتيح',
         'التأكد من أن كل مفتاح فريد (توليد 100 مفتاح)',
         '✅ نجاح - 100 مفتاح فريد'),
        ('اختبار 2.1.3: تشفير وفك تشفير الملفات',
         'التحقق من صحة تشفير الملفات',
         '✅ نجاح - المحتوى المسترجع مطابق للأصلي'),
        ('اختبار 2.1.4: Property-Based Testing للملفات',
         'اختبار تشفير ملفات بأحجام مختلفة (1-50KB) - 30 حالة',
         '✅ نجاح'),
        ('اختبار 2.1.5: تشفير مع مفتاح جديد',
         'التحقق من وظيفة encrypt_file_with_new_key',
         '✅ نجاح'),
        ('اختبار 2.1.6: رفض المفتاح الخاطئ للملفات',
         'التأكد من حماية الملفات من الوصول غير المصرح',
         '✅ نجاح - ValueError عند استخدام مفتاح خاطئ'),
    ]
    
    for title, goal, result in tests_2:
        doc.add_heading(title, level=3)
        p = doc.add_paragraph()
        p.add_run('الهدف: ').bold = True
        p.add_run(f'{goal}\n')
        p.add_run('النتيجة: ').bold = True
        p.add_run(result)
    
    doc.add_page_break()
    
    # القسم 3: PQ-X3DH
    doc.add_heading('3. اختبارات PQ-X3DH (تبادل المفاتيح الكمي)', level=1)
    
    doc.add_heading('3.1 البروتوكول المستخدم', level=2)
    p = doc.add_paragraph()
    p.add_run('• X25519: ').bold = True
    p.add_run('تبادل مفاتيح Diffie-Hellman على منحنى إهليلجي\n')
    p.add_run('• Kyber512: ').bold = True
    p.add_run('خوارزمية تشفير مقاومة للحوسبة الكمية\n')
    p.add_run('• HKDF-SHA256: ').bold = True
    p.add_run('اشتقاق المفاتيح')
    
    doc.add_heading('3.2 الاختبارات المنفذة', level=2)
    
    doc.add_heading('اختبار 3.2.1: اتفاق المفاتيح بين طرفين', level=3)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التحقق من أن Alice و Bob يصلان لنفس المفتاح الجذري\n')
    p.add_run('الطريقة:\n').bold = True
    p.add_run('  1. Alice تولد مفاتيحها (IK, EK)\n')
    p.add_run('  2. Bob يولد مفاتيحه (IK, SPK, OPK, Kyber)\n')
    p.add_run('  3. Alice تبدأ الجلسة (pqx3dh_initiate)\n')
    p.add_run('  4. Bob يستجيب (pqx3dh_respond)\n')
    p.add_run('  5. مقارنة المفاتيح الجذرية\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - المفاتيح متطابقة (32 بايت)')
    
    doc.add_heading('اختبار 3.2.2: PQ-X3DH بدون OPK', level=3)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التحقق من عمل البروتوكول بدون One-Time Pre-Key\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - البروتوكول يعمل بدون OPK')
    
    doc.add_heading('اختبار 3.2.3: جلسات مختلفة تنتج مفاتيح مختلفة', level=3)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التأكد من أن كل جلسة لها مفتاح فريد\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - مفاتيح Ephemeral جديدة لكل جلسة')
    
    doc.add_page_break()
    
    # القسم 4: Double Ratchet
    doc.add_heading('4. اختبارات Double Ratchet (تشفير الرسائل)', level=1)
    
    doc.add_heading('4.1 البروتوكول المستخدم', level=2)
    p = doc.add_paragraph()
    p.add_run('• Signal Protocol Double Ratchet\n')
    p.add_run('• KDF Chain: HKDF-SHA256\n')
    p.add_run('• تشفير الرسائل: XChaCha20-Poly1305')
    
    doc.add_heading('4.2 الاختبارات المنفذة', level=2)
    
    tests_4 = [
        ('اختبار 4.2.1: تبادل رسالة واحدة',
         'التحقق من تشفير وفك تشفير رسالة بين طرفين',
         '✅ نجاح - الرسالة "Hello Bob! مرحبا بوب" تم تشفيرها وفكها بنجاح'),
        ('اختبار 4.2.2: رسائل ثنائية الاتجاه',
         'التحقق من تبادل الرسائل في الاتجاهين (Alice↔Bob)',
         '✅ نجاح - جميع الرسائل وصلت بشكل صحيح'),
        ('اختبار 4.2.3: رسائل متعددة في نفس الاتجاه',
         'إرسال 10 رسائل متتالية من Alice إلى Bob',
         '✅ نجاح - جميع الرسائل وصلت بالترتيب الصحيح'),
        ('اختبار 4.2.4: رسائل خارج الترتيب',
         'معالجة الرسائل التي تصل بترتيب مختلف (3, 1, 2)',
         '✅ نجاح - يدعم الشبكات غير الموثوقة'),
        ('اختبار 4.2.5: حفظ واستعادة حالة Ratchet',
         'حفظ الجلسة (to_dict) واستئنافها (from_dict)',
         '✅ نجاح - الجلسة استُؤنفت بدون مشاكل'),
        ('اختبار 4.2.6: Property-Based Testing للرسائل',
         '20 حالة اختبار بأحجام 1-1000 بايت',
         '✅ نجاح'),
    ]
    
    for title, goal, result in tests_4:
        doc.add_heading(title, level=3)
        p = doc.add_paragraph()
        p.add_run('الهدف: ').bold = True
        p.add_run(f'{goal}\n')
        p.add_run('النتيجة: ').bold = True
        p.add_run(result)
    
    doc.add_page_break()
    
    # القسم 5: خصائص الأمان
    doc.add_heading('5. اختبارات خصائص الأمان', level=1)
    
    doc.add_heading('5.1 السرية الأمامية (Forward Secrecy)', level=2)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التأكد من أن اختراق مفتاح جلسة لا يكشف الرسائل السابقة\n')
    p.add_run('الطريقة:\n').bold = True
    p.add_run('  1. إنشاء جلسة 1 → مفتاح RK1\n')
    p.add_run('  2. إنشاء جلسة 2 → مفتاح RK2\n')
    p.add_run('  3. التحقق من أن RK1 ≠ RK2\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - كل جلسة لها مفاتيح Ephemeral فريدة')
    
    doc.add_heading('5.2 عشوائية المفاتيح', level=2)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التحقق من جودة توليد الأرقام العشوائية\n')
    p.add_run('الطريقة:\n').bold = True
    p.add_run('  1. توليد 1000 مفتاح\n')
    p.add_run('  2. التحقق من التفرد (1000 مفتاح فريد)\n')
    p.add_run('  3. التحقق من التوزيع (كل قيمة بايت 0-255 تظهر)\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - توزيع متساوي للقيم')
    
    doc.add_heading('5.3 عدم تمييز النصوص المشفرة', level=2)
    p = doc.add_paragraph()
    p.add_run('الهدف: ').bold = True
    p.add_run('التأكد من أن نفس النص يُشفر بشكل مختلف كل مرة\n')
    p.add_run('الطريقة: ').bold = True
    p.add_run('تشفير نفس النص 100 مرة بنفس المفتاح\n')
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ نجاح - 100 نص مشفر فريد\n')
    p.add_run('الأهمية: ').bold = True
    p.add_run('يمنع هجمات تحليل الأنماط')
    
    doc.add_page_break()
    
    # القسم 6: الخلاصة
    doc.add_heading('6. الخلاصة والتوصيات', level=1)
    
    doc.add_heading('6.1 نقاط القوة', level=2)
    strengths = [
        'تشفير قوي: XChaCha20-Poly1305 يوفر تشفير وتوثيق في آن واحد',
        'مقاومة الحوسبة الكمية: Kyber512 يحمي من الهجمات الكمية المستقبلية',
        'السرية الأمامية: كل جلسة لها مفاتيح فريدة',
        'كشف التلاعب: أي تعديل على البيانات يُكتشف فوراً',
        'دعم الرسائل خارج الترتيب: مرونة في الشبكات غير الموثوقة',
    ]
    for s in strengths:
        doc.add_paragraph(f'✅ {s}', style='List Bullet')
    
    doc.add_heading('6.2 المعايير المحققة', level=2)
    standards = [
        'NIST Post-Quantum: Kyber512 (FIPS 203)',
        'Signal Protocol: Double Ratchet',
        'IETF RFC 8439: ChaCha20-Poly1305',
    ]
    for s in standards:
        doc.add_paragraph(f'✅ {s}', style='List Bullet')
    
    doc.add_heading('6.3 التوصيات', level=2)
    recommendations = [
        'تحديث مكتبة liboqs لتتوافق مع liboqs-python',
        'إضافة اختبارات أداء للتشفير على ملفات كبيرة',
        'إضافة اختبارات تكامل مع واجهة المستخدم',
    ]
    for i, r in enumerate(recommendations, 1):
        doc.add_paragraph(f'{i}. {r}')
    
    doc.add_paragraph('─' * 50)
    
    # بيئة الاختبار
    doc.add_heading('7. بيئة الاختبار', level=1)
    env = doc.add_paragraph()
    env.add_run('نظام التشغيل: ').bold = True
    env.add_run('Kali Linux\n')
    env.add_run('Python: ').bold = True
    env.add_run('3.10.13\n')
    env.add_run('pytest: ').bold = True
    env.add_run('9.0.2\n')
    env.add_run('hypothesis: ').bold = True
    env.add_run('6.149.1\n')
    env.add_run('liboqs: ').bold = True
    env.add_run('0.13.1-dev\n')
    env.add_run('oqs-python: ').bold = True
    env.add_run('0.14.0')
    
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    
    # الخاتمة
    conclusion = doc.add_paragraph()
    conclusion.add_run('تم إعداد هذا التقرير آلياً بواسطة نظام الاختبار\n').italic = True
    conclusion.add_run('جميع الاختبارات نجحت ✅').bold = True
    conclusion.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # حفظ الملف
    doc.save('CRYPTO_TEST_REPORT.docx')
    print('✅ تم إنشاء الملف: CRYPTO_TEST_REPORT.docx')

if __name__ == '__main__':
    create_report()
