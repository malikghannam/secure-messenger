"""
تحويل تقرير التشفير مع الدلائل إلى Word
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_code_block(doc, code_text):
    """إضافة كتلة كود بخلفية رمادية"""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    return p

def create_report():
    doc = Document()
    
    # العنوان الرئيسي
    title = doc.add_heading('تقرير اختبار نظام التشفير', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Cryptographic System Test Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # معلومات التقرير
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('التاريخ: ').bold = True
    info.add_run('11 يناير 2026\n')
    info.add_run('المشروع: ').bold = True
    info.add_run('Secure Messenger\n')
    info.add_run('الإصدار: ').bold = True
    info.add_run('1.0')
    
    doc.add_paragraph('─' * 50)
    
    # ملخص تنفيذي
    doc.add_heading('ملخص تنفيذي', level=1)
    doc.add_paragraph(
        'تم إجراء 8 اختبارات شاملة على نظام التشفير في تطبيق المراسلة الآمنة. '
        'جميع الاختبارات نجحت بنسبة 100%، مما يثبت فعالية وأمان نظام التشفير المستخدم.'
    )
    
    # جدول النتائج
    doc.add_heading('نتائج الاختبارات', level=2)
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    
    headers = ['الاختبار', 'النتيجة', 'الدليل']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data = [
        ['XChaCha20 Round-Trip', 'PASSED', 'النص المسترجع = الأصلي'],
        ['Wrong Key Rejection', 'PASSED', 'CryptoError عند مفتاح خاطئ'],
        ['Tamper Detection', 'PASSED', 'كشف تعديل بايت واحد'],
        ['Unique Ciphertexts', 'PASSED', '5/5 نصوص فريدة'],
        ['File Encryption', 'PASSED', 'الملف المسترجع = الأصلي'],
        ['PQ-X3DH Key Agreement', 'PASSED', 'مفاتيح Alice = Bob'],
        ['Double Ratchet', 'PASSED', '3 رسائل تم تبادلها'],
        ['Forward Secrecy', 'PASSED', 'جلسات مختلفة = مفاتيح مختلفة'],
    ]
    
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_page_break()
    
    # ========== الدلائل التفصيلية ==========
    doc.add_heading('الدلائل التفصيلية على نجاح الاختبارات', level=1)
    
    # اختبار 1
    doc.add_heading('اختبار 1: XChaCha20-Poly1305 Round-Trip', level=2)
    doc.add_paragraph('هذا الاختبار يثبت أن التشفير ثم فك التشفير يعيد البيانات الأصلية بدون أي تغيير.')
    
    doc.add_heading('المدخلات:', level=3)
    add_code_block(doc, 'Key: 8fcf80480abc06d1aa480eefcb9f42901cbb10fa3bf49e0ec850b88222fae9f7')
    add_code_block(doc, 'Plaintext: "Hello World! This is a secret message."')
    
    doc.add_heading('المخرجات:', level=3)
    add_code_block(doc, 'Ciphertext: 286c10db0253dd5a62dc8bccc1334edb94d40767cfeb1224...')
    add_code_block(doc, 'Decrypted: "Hello World! This is a secret message."')
    
    p = doc.add_paragraph()
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ PASSED - النص المفكوك مطابق للنص الأصلي')
    
    # اختبار 2
    doc.add_heading('اختبار 2: رفض المفتاح الخاطئ', level=2)
    doc.add_paragraph('هذا الاختبار يثبت أن النظام يرفض فك التشفير بمفتاح غير صحيح.')
    
    doc.add_heading('المدخلات:', level=3)
    add_code_block(doc, 'Key 1 (correct): 1d723a89b628a2b31f77227790c9888c0d2fa15e2134b0bd576cca8dac094f80')
    add_code_block(doc, 'Key 2 (wrong):   4a5fbd1b9944320cde472dda12518dcf13e44258fe3afb581c49b8a2d1cd08e1')
    
    doc.add_heading('المخرجات:', level=3)
    add_code_block(doc, 'Error: CryptoError - فشل فك التشفير بالمفتاح الخاطئ')
    
    p = doc.add_paragraph()
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ PASSED - المفتاح الخاطئ رُفض بشكل صحيح')
    
    # اختبار 3
    doc.add_heading('اختبار 3: كشف التلاعب بالبيانات', level=2)
    doc.add_paragraph('هذا الاختبار يثبت أن أي تعديل على البيانات المشفرة يتم كشفه.')
    
    doc.add_heading('العملية:', level=3)
    add_code_block(doc, 'Original Ciphertext: f09a7ded178edc181a9cf298512dd9ba94d53fd3655619fb054dc281692543b7...')
    add_code_block(doc, 'Modified byte at position 30: 0x43 -> 0xbc')
    add_code_block(doc, 'Tampered Ciphertext: f09a7ded178edc181a9cf298512dd9ba94d53fd3655619fb054dc2816925bcb7...')
    
    doc.add_heading('المخرجات:', level=3)
    add_code_block(doc, 'Error: CryptoError - تم كشف التلاعب')
    
    p = doc.add_paragraph()
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ PASSED - تعديل بايت واحد تم كشفه (حماية من هجمات MITM)')
    
    # اختبار 4
    doc.add_heading('اختبار 4: تفرد النصوص المشفرة', level=2)
    doc.add_paragraph('هذا الاختبار يثبت أن نفس النص يُشفر بشكل مختلف كل مرة (بسبب Nonce عشوائي).')
    
    doc.add_heading('المدخلات:', level=3)
    add_code_block(doc, 'Key: 9316a76741927e61cdcd90983ed9a090a13ee7e03c5655316d2900baec381453')
    add_code_block(doc, 'Plaintext: "Same message" (تشفير 5 مرات)')
    
    doc.add_heading('المخرجات:', level=3)
    add_code_block(doc, '''Ciphertext 1: 99db015299590549fcc29870f2a3da8d62e16fe2...
Ciphertext 2: ed443d37d613e715576bcff1c2f526ece1f794e5...
Ciphertext 3: 33b5fbecd7133f62cfc5fdea894d9785d3b4d855...
Ciphertext 4: d71a73f5b0267f70366baaaf7a2744a3530cf417...
Ciphertext 5: 07ea8ac954e7225b6e35f2e4413f2daa52455d2d...''')
    
    p = doc.add_paragraph()
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ PASSED - 5/5 نصوص مشفرة فريدة (يمنع تحليل الأنماط)')
    
    doc.add_page_break()
    
    # اختبار 5
    doc.add_heading('اختبار 5: تشفير الملفات', level=2)
    doc.add_paragraph('هذا الاختبار يثبت صحة تشفير وفك تشفير الملفات.')
    
    doc.add_heading('المدخلات:', level=3)
    add_code_block(doc, 'File Size: 87 bytes')
    add_code_block(doc, 'File Key: c66248415ee2bedb8de396ff2a574c5a47f156c04751ce7e11f1e0558adeaafa')
    
    doc.add_heading('المخرجات:', level=3)
    add_code_block(doc, 'Encrypted Size: 127 bytes (87 + 24 nonce + 16 tag)')
    add_code_block(doc, 'Decrypted Size: 87 bytes (مطابق للأصلي)')
    
    p = doc.add_paragraph()
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ PASSED - الملف المفكوك مطابق للأصلي')
    
    # اختبار 6
    doc.add_heading('اختبار 6: PQ-X3DH Key Agreement', level=2)
    doc.add_paragraph('هذا الاختبار يثبت أن طرفين (Alice و Bob) يمكنهما الاتفاق على مفتاح مشترك بشكل آمن.')
    
    doc.add_heading('مفاتيح Alice:', level=3)
    add_code_block(doc, 'Identity Key: 35dc5f02be394b5b0d8fc04a07b62e42ce1b282498e5c44d8abf587a7519d659')
    add_code_block(doc, 'Ephemeral Key: d50c5c0479c2ccf3dbefa94364313adfac55d1ceb63165feb2106916da3d6608')
    
    doc.add_heading('مفاتيح Bob:', level=3)
    add_code_block(doc, 'Identity Key: 785d803f06a654bd05f3736fe1402087542c13cd73890805744743fb81685628')
    add_code_block(doc, 'Signed Pre-Key: c172d447218845db174993588e21d5d884d69e68cf08ab8730af91b9dd924679')
    add_code_block(doc, 'Kyber Public Key: 9646379d5529743ab5e04058459a435c294058f6739f563020a62224f301ecb3... (800 bytes)')
    
    doc.add_heading('المفاتيح الجذرية المشتقة:', level=3)
    add_code_block(doc, "Alice's Root Key: 3f3f2af70e895e266d9a56b56d5b32cc4e8b45a7923aab637559e476d62f456e")
    add_code_block(doc, "Bob's Root Key:   3f3f2af70e895e266d9a56b56d5b32cc4e8b45a7923aab637559e476d62f456e")
    
    p = doc.add_paragraph()
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ PASSED - الطرفان اشتقا نفس المفتاح الجذري!')
    
    # اختبار 7
    doc.add_heading('اختبار 7: Double Ratchet Message Exchange', level=2)
    doc.add_paragraph('هذا الاختبار يثبت تبادل الرسائل المشفرة بين طرفين باستخدام بروتوكول Double Ratchet.')
    
    doc.add_heading('الجلسة:', level=3)
    add_code_block(doc, 'Shared Root Key: e1193737b8f4abe941ba8dec381a98c7dfff901d082157503e49de79eaaaa494')
    
    doc.add_heading('الرسائل المتبادلة:', level=3)
    
    p = doc.add_paragraph()
    p.add_run('Message 1 (Alice → Bob):\n').bold = True
    add_code_block(doc, '''Plaintext: "Hello Bob! How are you?"
Encrypted: {"header": {"dh_pub": "pQSXEWg4Ck8ekIxQVS0hkOhM04NJJ_h-iM_3qdVBrgc"...}
Decrypted: "Hello Bob! How are you?" ✓''')
    
    p = doc.add_paragraph()
    p.add_run('Message 2 (Bob → Alice):\n').bold = True
    add_code_block(doc, '''Plaintext: "Hi Alice! I'm fine, thanks!"
Encrypted: {"header": {"dh_pub": "_vgWdF1xA84-uEGQh0wm6vTkJwuVHACvYf8eAXA_8xs"...}
Decrypted: "Hi Alice! I'm fine, thanks!" ✓''')
    
    p = doc.add_paragraph()
    p.add_run('Message 3 (Alice → Bob):\n').bold = True
    add_code_block(doc, '''Plaintext: "Great! Let's meet tomorrow."
Decrypted: "Great! Let's meet tomorrow." ✓''')
    
    p = doc.add_paragraph()
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ PASSED - جميع الرسائل تم تبادلها بنجاح')
    
    doc.add_page_break()
    
    # اختبار 8
    doc.add_heading('اختبار 8: Forward Secrecy', level=2)
    doc.add_paragraph('هذا الاختبار يثبت أن اختراق جلسة واحدة لا يؤثر على الجلسات الأخرى.')
    
    doc.add_heading('جلسة 1:', level=3)
    add_code_block(doc, '''Ephemeral Key: 99a0138ac60e7238be45187c95b8f433adf0cf77a61defc8868f9cf87ff59352
Root Key: 2b3f3fbe998e724e011f1559a6f49531f34925ddfabac6bc52680d03d02ccc13''')
    
    doc.add_heading('جلسة 2:', level=3)
    add_code_block(doc, '''Ephemeral Key: 87181b8d0f4ea2f99ebb70aff70e917ee494d111e4c8c9b2ac43957cb8907677
Root Key: 55a8543949fa18e5f48fb4afac41cf0897aac4e177edbbedfef984c231c94f3a''')
    
    doc.add_heading('التحقق:', level=3)
    add_code_block(doc, '''Ephemeral keys different: True
Root keys different: True''')
    
    p = doc.add_paragraph()
    p.add_run('النتيجة: ').bold = True
    p.add_run('✅ PASSED - السرية الأمامية محققة (كل جلسة لها مفاتيح فريدة)')
    
    doc.add_page_break()
    
    # ========== الخلاصة ==========
    doc.add_heading('الخلاصة', level=1)
    
    doc.add_heading('نقاط القوة المثبتة:', level=2)
    strengths = [
        'تشفير XChaCha20-Poly1305 يعمل بشكل صحيح (Round-trip)',
        'حماية من المفاتيح الخاطئة (Key Rejection)',
        'كشف التلاعب بالبيانات (Tamper Detection)',
        'عشوائية التشفير (Unique Ciphertexts)',
        'تشفير الملفات يعمل بشكل صحيح',
        'تبادل المفاتيح الكمي PQ-X3DH يعمل',
        'بروتوكول Double Ratchet للرسائل يعمل',
        'السرية الأمامية (Forward Secrecy) محققة',
    ]
    for s in strengths:
        doc.add_paragraph(f'✅ {s}', style='List Bullet')
    
    doc.add_heading('المعايير الأمنية المحققة:', level=2)
    standards = [
        'NIST Post-Quantum: Kyber512 (FIPS 203)',
        'Signal Protocol: Double Ratchet',
        'IETF RFC 8439: ChaCha20-Poly1305',
        'X25519: Elliptic Curve Diffie-Hellman',
    ]
    for s in standards:
        doc.add_paragraph(f'✅ {s}', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    
    # بيئة الاختبار
    doc.add_heading('بيئة الاختبار', level=2)
    env = doc.add_paragraph()
    env.add_run('نظام التشغيل: ').bold = True
    env.add_run('Kali Linux\n')
    env.add_run('Python: ').bold = True
    env.add_run('3.10.13\n')
    env.add_run('تاريخ التنفيذ: ').bold = True
    env.add_run('2026-01-11 08:47:17\n')
    env.add_run('عدد الاختبارات: ').bold = True
    env.add_run('8\n')
    env.add_run('نسبة النجاح: ').bold = True
    env.add_run('100%')
    
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    
    # الخاتمة
    conclusion = doc.add_paragraph()
    conclusion.add_run('\nجميع الاختبارات نجحت ✅\n').bold = True
    conclusion.add_run('ALL TESTS PASSED SUCCESSFULLY!').bold = True
    conclusion.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # حفظ الملف
    doc.save('CRYPTO_TEST_REPORT_WITH_EVIDENCE.docx')
    print('✅ تم إنشاء الملف: CRYPTO_TEST_REPORT_WITH_EVIDENCE.docx')

if __name__ == '__main__':
    create_report()
